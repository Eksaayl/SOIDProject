from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('Bulleted List Example', level=1)

items = [
    "First bullet point",
    "Second bullet point",
    "Third bullet point"
]

for item in items:
    para = doc.add_paragraph(f"• {item}")
    run = para.runs[0]
    run.font.name = "Palatino Linotype"
    run.font.size = Pt(12)  

doc.save("templates.docx")
print("Document created: templates.docx")