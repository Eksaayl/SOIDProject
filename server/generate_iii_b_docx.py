from docx import Document
from docx.shared import Pt

def create_iii_b_docx(projects, output_path):
    doc = Document()
    for project in projects:
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        fields = [
            ('A.1 NAME/TITLE', project.get('name', '')),
            ('A.2 OBJECTIVES', project.get('objectives', '')),
            ('A.3 DURATION', project.get('duration', '')),
            ('A.4 DELIVERABLES', project.get('deliverables', [])),
            ('A.5 LEAD AGENCY', project.get('lead_agency', '')),
            ('A.6 IMPLEMENTING AGENCIES', project.get('implementing_agencies', '')),
        ]
        for i, (label, value) in enumerate(fields):
            table.cell(i, 0).text = label
            if label == 'A.4 DELIVERABLES':
                cell = table.cell(i, 1)
                # Remove all existing paragraphs (including the default empty one)
                for para in list(cell.paragraphs):
                    p = para._element
                    p.getparent().remove(p)
                if isinstance(value, str):
                    value = [v.strip() for v in value.split('\n') if v.strip()]
                for deliverable in value:
                    p = cell.add_paragraph(deliverable, style='List Bullet')
                    p.paragraph_format.left_indent = Pt(0)
                    cell.add_paragraph('')  # Add a blank line after each bullet
            else:
                table.cell(i, 1).text = value
        doc.add_paragraph()  # Space between tables
    doc.save(output_path)

# Example usage:
if __name__ == '__main__':
    projects = [
        {
            'name': 'National ICT Household Survey Data Processing and Management System (NICTHS)',
            'objectives': 'To be used in the processing of survey results data on ICT for household and individual respondents.',
            'duration': '2024-2026 (yearly)',
            'deliverables': [
                'Online facility for monitoring of distribution, collection, and processing of survey questionnaires that will provide validation and generation of reports and tables on the data collected by the PSA Field Offices.'
            ],
            'lead_agency': 'Department of Information and Communications Technology (DICT)',
            'implementing_agencies': 'Philippine Statistics Authority (PSA)',
        },
    ]
    create_iii_b_docx(projects, 'iii_b_projects.docx') 