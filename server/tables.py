from docx import Document
from docx.shared import Pt, Inches

def create_iib_docx(systems, output_path, template_path=None):
    if template_path is None:
        template_path = r'C:/Users/User/Documents/SOIDProject/assets/II_b.docx'
    doc = Document(template_path)
    for sys in systems:
        table = doc.add_table(rows=8, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = 'NAME OF INFORMATION SYSTEM/ SUB-SYSTEM'
        table.cell(0, 2).text = sys.get('name_of_system', '')

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 0).text = 'DESCRIPTION'
        cell = table.cell(1, 2)
        for para in list(cell.paragraphs):
            p = para._element
            p.getparent().remove(p)
        value = sys.get('description', [])
        if isinstance(value, str):
            value = [v.strip() for v in value.split('\n') if v.strip()]
        for item in value:
            p = cell.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Pt(0)
            cell.add_paragraph('')

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 0).text = 'STATUS'
        table.cell(2, 2).text = sys.get('status', '')

        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(3, 0).text = 'DEVELOPMENT STRATEGY'
        table.cell(3, 2).text = sys.get('development_strategy', '')

        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(4, 0).text = 'COMPUTING SCHEME'
        table.cell(4, 2).text = sys.get('computing_scheme', '')

        table.cell(5, 0).merge(table.cell(6, 0))
        table.cell(5, 0).text = 'USERS'
        table.cell(5, 1).text = 'INTERNAL'
        internal_users = sys.get('users_internal', [])
        if isinstance(internal_users, str):
            internal_users = [internal_users]
        table.cell(5, 2).text = '\n'.join(internal_users)

        table.cell(6, 1).text = 'EXTERNAL'
        table.cell(6, 2).text = sys.get('users_external', '')

        table.cell(7, 0).merge(table.cell(7, 1))
        table.cell(7, 0).text = 'OWNER'
        table.cell(7, 2).text = sys.get('owner', '')

        for i in range(8):
            cell = table.cell(i, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Palatino Linotype'
                    run.font.size = Pt(12)
            if i == 0:
                cell = table.cell(i, 2)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Palatino Linotype'
                        run.font.size = Pt(12)

        col1_width = Inches(0.75)   
        col2_width = Inches(1)    
        col3_width = Inches(7.19)   
        for i, row in enumerate(table.rows):
            if i == 5:
                row.cells[0].width = Inches(0.5)
            else:
                row.cells[0].width = col1_width
            row.cells[1].width = col2_width
            row.cells[2].width = col3_width

        doc.add_paragraph()
    doc.save(output_path)

def create_iic_docx(databases, output_path, template_path=None):
    if template_path is None:
        template_path = r'C:/Users/User/Documents/SOIDProject/assets/II_c.docx'
    doc = Document(template_path)
    for db in databases:
        table = doc.add_table(rows=8, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = 'NAME OF DATABASE'
        table.cell(0, 2).text = db.get('name_of_database', '')

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 0).text = 'GENERAL CONTENTS/DESCRIPTION'
        cell = table.cell(1, 2)
        for para in list(cell.paragraphs):
            p = para._element
            p.getparent().remove(p)
        value = db.get('general_contents', [])
        if isinstance(value, str):
            value = [v.strip() for v in value.split('\n') if v.strip()]
        for item in value:
            p = cell.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Pt(0)
            cell.add_paragraph('')

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 0).text = 'STATUS'
        table.cell(2, 2).text = db.get('status', '')

        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(3, 0).text = 'INFORMATION SYSTEMS SERVED'
        table.cell(3, 2).text = db.get('info_systems_served', '')

        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(4, 0).text = 'DATA ARCHIVING/STORAGE MEDIA'
        table.cell(4, 2).text = db.get('data_archiving', '')

        table.cell(5, 0).merge(table.cell(6, 0))
        table.cell(5, 0).text = 'USERS'
        table.cell(5, 1).text = 'INTERNAL'
        internal_users = db.get('users_internal', [])
        if isinstance(internal_users, str):
            internal_users = [internal_users]
        table.cell(5, 2).text = '\n'.join(internal_users)

        table.cell(6, 1).text = 'EXTERNAL'
        table.cell(6, 2).text = db.get('users_external', '')

        table.cell(7, 0).merge(table.cell(7, 1))
        table.cell(7, 0).text = 'OWNER'
        table.cell(7, 2).text = db.get('owner', '')

        for i in range(8):
            cell = table.cell(i, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Palatino Linotype'
                    run.font.size = Pt(12)
            if i == 0:
                cell = table.cell(i, 2)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Palatino Linotype'
                        run.font.size = Pt(12)
        col1_width = Inches(0.75)   
        col2_width = Inches(1)    
        col3_width = Inches(7.19)   
        for i, row in enumerate(table.rows):
            if i == 5:
                row.cells[0].width = Inches(0.5)
            else:
                row.cells[0].width = col1_width
            row.cells[1].width = col2_width
            row.cells[2].width = col3_width

        doc.add_paragraph()
    doc.save(output_path)

def create_iii_a_docx(projects, output_path, template_path=None):
    if template_path is None:
        template_path = r'C:/Users/User/Documents/SOIDProject/assets/III_a.docx'
    doc = Document(template_path)
    for idx, project in enumerate(projects):
        rank_para = doc.add_paragraph()
        rank_run = rank_para.add_run(f'RANK {idx + 1}')
        rank_run.bold = True
        rank_run.font.name = 'Palatino Linotype'
        rank_run.font.size = Pt(12)
        rank_para.paragraph_format.space_after = Pt(0)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        fields = [
            ('A.1 NAME/TITLE', project.get('name', '')),
            ('A.2 OBJECTIVES', project.get('objectives', '')),
            ('A.3 DURATION', project.get('duration', '')),
            ('A.4 DELIVERABLES', project.get('deliverables', [])),
        ]
        for i, (label, value) in enumerate(fields):
            table.cell(i, 0).text = label
            if label == 'A.4 DELIVERABLES':
                cell = table.cell(i, 1)
                for para in list(cell.paragraphs):
                    p = para._element
                    p.getparent().remove(p)
                if isinstance(value, str):
                    value = [v.strip() for v in value.split('\n') if v.strip()]
                for deliverable in value:
                    p = cell.add_paragraph(deliverable, style='List Bullet')
                    p.paragraph_format.left_indent = Pt(0)
                    cell.add_paragraph('')
            else:
                cell = table.cell(i, 1)
                cell.text = value

        for i in range(4):
            cell = table.cell(i, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Palatino Linotype'
                    run.font.size = Pt(12)
            if i == 0:
                cell = table.cell(i, 1)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Palatino Linotype'
                        run.font.size = Pt(12)
        for row in table.rows:
            row.cells[0].width = Inches(1.77165)
            row.cells[1].width = Inches(7.19)
        doc.add_paragraph()
    doc.save(output_path)

def create_iii_b_docx(projects, output_path, template_path=None):
    if template_path is None:
        template_path = r'C:/Users/User/Documents/SOIDProject/assets/III_b.docx'
    doc = Document(template_path)
    for project in projects:
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        fields = [
            ('A.1 NAME/TITLE', project.get('name', '')),
            ('A.2 OBJECTIVES', project.get('objectives', '')),
            ('A.3 DURATION', project.get('duration', '')),
            ('A.4 DELIVERABLES', project.get('deliverables', [])),
            ('A.5 LEAD AGENCY', project.get('lead_agency', '')),
            ('A.6 IMPLEMENTING AGENCIES', project.get('implementing_agencies', '')),
        ]
        for i, (label, value) in enumerate(fields):
            table.cell(i, 0).text = label
            if label == 'A.4 DELIVERABLES':
                cell = table.cell(i, 1)
                for para in list(cell.paragraphs):
                    p = para._element
                    p.getparent().remove(p)
                if isinstance(value, str):
                    value = [v.strip() for v in value.split('\n') if v.strip()]
                for deliverable in value:
                    p = cell.add_paragraph(deliverable, style='List Bullet')
                    p.paragraph_format.left_indent = Pt(0)
                    cell.add_paragraph('')
            else:
                cell = table.cell(i, 1)
                cell.text = value

        for i in range(6):
            cell = table.cell(i, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Palatino Linotype'
                    run.font.size = Pt(12)
            if i == 0:
                cell = table.cell(i, 1)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.name = 'Palatino Linotype'
                        run.font.size = Pt(12)
        for row in table.rows:
            row.cells[0].width = Inches(1.77165)
            row.cells[1].width = Inches(7.19)
        doc.add_paragraph()
    doc.save(output_path)

def create_iiic_logframe_table(data, output_path, doc=None):
    """
    data: dict with keys 'intermediate' (list of dicts), 'immediate', 'outputs'
    doc: optional, an existing Document object to add the table to
    """
    try:
        if doc is None:
            doc = Document()
        
        # Handle the data structure - data should be a single logframe
        intermediate_rows = data.get("intermediate", [])
        immediate_rows = data.get("immediate", [])
        outputs_rows = data.get("outputs", [])

        if not isinstance(intermediate_rows, list):
            intermediate_rows = [intermediate_rows] if intermediate_rows else []
        if not isinstance(immediate_rows, list):
            immediate_rows = [immediate_rows] if immediate_rows else []
        if not isinstance(outputs_rows, list):
            outputs_rows = [outputs_rows] if outputs_rows else []

        total_rows = 1 + len(intermediate_rows) + len(immediate_rows) + len(outputs_rows)
        table = doc.add_table(rows=total_rows, cols=6)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False

        col_widths = [2.11, 1.97, 0.98, 1.28, 1.05, 1.55]
        for col_idx, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = Inches(width)

        headers = [
            "Hierarchy of targeted results",
            "Objectively verifiable indicators (OVI)",
            "Baseline data",
            "Targets",
            "Data collection methods",
            "Responsibility to collect data"
        ]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)

        sections = [
            ('Intermediate Outcome:', intermediate_rows),
            ('Immediate Outcome:', immediate_rows),
            ('Outputs:', outputs_rows),
        ]
        row_idx = 1
        for label, rows in sections:
            for idx, row_data in enumerate(rows):
                cell = table.cell(row_idx, 0)
                if idx == 0:
                    cell.text = label
                    cell.add_paragraph("")
                else:
                    cell.text = ""
                user_value = row_data.get("hierarchy", "")
                if user_value:
                    for line in user_value.splitlines():
                        if line.strip():
                            cell.add_paragraph(line.strip())
                for col, col_key in enumerate(["ovi", "baseline", "targets", "methods", "responsibility"], start=1):
                    value = row_data.get(col_key, "")
                    cell = table.cell(row_idx, col)
                    cell.text = ""
                    if idx == 0:
                        cell.add_paragraph("")
                    if value:
                        for line in value.splitlines():
                            if line.strip():
                                cell.add_paragraph(line.strip())
                row_idx += 1

        # Only save if no existing document was passed
        if doc is None and output_path:
            doc.save(output_path)
    
    except Exception as e:
        print(f"Error in create_iiic_logframe_table: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

    