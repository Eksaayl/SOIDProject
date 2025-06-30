import os
import io
import base64
import tempfile
import logging
import traceback
from typing import List, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Inches, Pt, RGBColor
import uvicorn
import uuid
from server.tables import create_iii_b_docx, create_iii_a_docx, create_iiic_logframe_table
import argparse
import mammoth
from server.part_ib_handler import generate_part_ib_docx, process_part_ib_data
from datetime import datetime

logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def append_docx_with_section_break(master, doc_to_append):
    while master.paragraphs and not master.paragraphs[-1].text.strip():
        p = master.paragraphs[-1]._element
        p.getparent().remove(p)

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = master.paragraphs[-1]._element
    pPr = p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    type_ = OxmlElement('w:type')
    type_.set(qn('w:val'), 'nextPage')
    sectPr.append(type_)
    pPr.append(sectPr)

    if hasattr(doc_to_append, 'part'):
        if hasattr(doc_to_append.part, 'styles_part'):
            if not hasattr(master.part, 'styles_part'):
                master.part.styles_part = doc_to_append.part.styles_part
            else:
                for style in doc_to_append.part.styles_part.element:
                    existing_style = master.part.styles_part.element.find(
                        f'.//w:style[@w:styleId="{style.get(qn("w:styleId"))}"]',
                        {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    )
                    if existing_style is None:
                        master.part.styles_part.element.append(style)

        if hasattr(doc_to_append.part, 'numbering_part'):
            if not hasattr(master.part, 'numbering_part'):
                master.part.numbering_part = doc_to_append.part.numbering_part
            else:
                for num in doc_to_append.part.numbering_part.element:
                    master.part.numbering_part.element.append(num)

    for element in doc_to_append.element.body:
        master.element.body.append(element)

    for paragraph in master.paragraphs:
        if hasattr(paragraph, '_element'):
            p = paragraph._element
            if hasattr(paragraph, 'style') and paragraph.style:
                pPr = p.get_or_add_pPr()
                pStyle = OxmlElement('w:pStyle')
                pStyle.set(qn('w:val'), paragraph.style.name)
                pPr.append(pStyle)

def insert_bullet_list_at_placeholder(doc_path, output_path, placeholder, items):
    from docx import Document
    doc = Document(doc_path)
    for para in doc.paragraphs:
        if placeholder in para.text:
            parent = para._element.getparent()
            idx = list(parent).index(para._element)
            p = para._element
            p.getparent().remove(p)
            for item in items:
                new_para = doc.add_paragraph(item, style='List Bullet')
                parent.insert(idx, new_para._element)
                idx += 1
            break
    doc.save(output_path)

async def generate_docx_II_a(template_file, images):
    try:
        doc = Document(template_file)
        sorted_images = sorted(images, key=lambda x: x.filename)
        image_map = {
            'ISI': None,
            'ISII': None,
            'ISIII': None,
        }
        
        for i, img in enumerate(sorted_images):
            if i == 0:
                image_map['ISI'] = img
            elif i == 1:
                image_map['ISII'] = img
            elif i == 2:
                image_map['ISIII'] = img
        
        for paragraph in doc.paragraphs:
            if '{ISI}' in paragraph.text:
                if image_map['ISI']:
                    img_content = await image_map['ISI'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(8.52), height=Inches(5.69))
            elif '{ISII}' in paragraph.text:
                if image_map['ISII']:
                    img_content = await image_map['ISII'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(8.52), height=Inches(5.69))
            elif '{ISIII}' in paragraph.text:
                if image_map['ISIII']:
                    img_content = await image_map['ISIII'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(8.52), height=Inches(5.69))
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Part II.A document: {str(e)}")

async def generate_docx_II_d(template_file, images):
    try:
        doc = Document(r'C:/Users/User/Documents/SOIDProject/assets/II_d.docx')
        sorted_images = sorted(images, key=lambda x: x.filename)
        image_map = {
            'NLC': None,
            'PNL': None,
        }
        
        for i, img in enumerate(sorted_images):
            if i == 0:
                image_map['NLC'] = img
            elif i == 1:
                image_map['PNL'] = img
        
        for paragraph in doc.paragraphs:
            if '{NLC}' in paragraph.text:
                if image_map['NLC']:
                    img_content = await image_map['NLC'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(6))
            elif '{PNL}' in paragraph.text:
                if image_map['PNL']:
                    img_content = await image_map['PNL'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(6))
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Part II.D document: {str(e)}")

async def generate_docx_IV_b(template_file, images):
    try:
        doc = Document(template_file)
        sorted_images = sorted(images, key=lambda x: x.filename)
        image_map = {
            'Existing': None,
            'Proposed': None,
            'Placement': None,
        }
        
        for i, img in enumerate(sorted_images):
            if i == 0:
                image_map['Existing'] = img
            elif i == 1:
                image_map['Proposed'] = img
            elif i == 2:
                image_map['Placement'] = img
        
        for paragraph in doc.paragraphs:
            if '{Existing}' in paragraph.text:
                if image_map['Existing']:
                    img_content = await image_map['Existing'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(6))
            elif '{Proposed}' in paragraph.text:
                if image_map['Proposed']:
                    img_content = await image_map['Proposed'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(6))
            elif '{Placement}' in paragraph.text:
                if image_map['Placement']:
                    img_content = await image_map['Placement'].read()
                    img_stream = io.BytesIO(img_content)
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(6))
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Part IV.B document: {str(e)}")

@app.post("/generate-docx")
async def generate_docx(
    template: UploadFile = File(...),
    images: List[UploadFile] = File(...),
    request: Request = None,
):
    template_file = None
    try:
        if not template.filename:
            raise HTTPException(status_code=400, detail="Template filename is required")
        
        # Get yearRange from headers
        year_range = request.headers.get('yearrange', '') if request else ''
        logger.debug(f"[GENERATE DOCX] Received yearRange: '{year_range}'")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        template_file = f"temp_template_{timestamp}.docx"
        
        try:
            with open(template_file, "wb") as f:
                f.write(await template.read())
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save template file: {str(e)}")
        
        if not images:
            raise HTTPException(status_code=400, detail="At least one image is required")
        
        try:
            if "II_a.docx" in template.filename:
                if len(images) != 3:
                    raise HTTPException(status_code=400, detail="Part II.A requires exactly 3 images")
                docx_bytes = await generate_docx_II_a(template_file, images)
            elif "II_d.docx" in template.filename:
                if len(images) != 2:
                    raise HTTPException(status_code=400, detail="Part II.D requires exactly 2 images")
                docx_bytes = await generate_docx_II_d(template_file, images)
            elif "IV_b.docx" in template.filename:
                if len(images) != 3:
                    raise HTTPException(status_code=400, detail="Part IV.B requires exactly 3 images")
                # Apply yearRange replacements for IV.B if provided
                if year_range:
                    temp_template_with_year = f"temp_iv_b_with_year_{timestamp}.docx"
                    template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'assets', 'IV_b.docx')
                    replacements = {"${yearRange}": year_range}
                    fill_placeholders_and_bullets(template_path, temp_template_with_year, replacements)
                    template_file = temp_template_with_year
                docx_bytes = await generate_docx_IV_b(template_file, images)
            else:
                raise HTTPException(status_code=400, detail=f"Invalid template file: {template.filename}")
        except HTTPException:
            raise
        except Exception as e:
            print("Exception in /generate-docx:", e)
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")
        
        if "IV_b.docx" in template.filename:
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': 'attachment; filename="document.docx"'}
            )
        else:
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except HTTPException:
        raise
    except Exception as e:
        print("Exception in /generate-docx:", e)
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")
    finally:
        if template_file and os.path.exists(template_file):
            try:
                os.remove(template_file)
            except Exception as e:
                print(f"Warning: Failed to remove temporary file {template_file}: {str(e)}")

@app.post("/merge-documents-part-i-all")
async def merge_documents_part_i_all(
    part_ia: UploadFile = File(...),
    part_ib: UploadFile = File(...),
    part_ic: UploadFile = File(...),
    part_id: UploadFile = File(...),
    part_ie: UploadFile = File(...),
):
    try:
        try:
            ia_bytes = io.BytesIO(await part_ia.read())
            ib_bytes = io.BytesIO(await part_ib.read())
            ic_bytes = io.BytesIO(await part_ic.read())
            id_bytes = io.BytesIO(await part_id.read())
            ie_bytes = io.BytesIO(await part_ie.read())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read uploaded files: {str(e)}")

        try:
            ia_doc = Document(ia_bytes)
            ib_doc = Document(ib_bytes)
            ic_doc = Document(ic_bytes)
            id_doc = Document(id_bytes)
            ie_doc = Document(ie_bytes)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX files: {str(e)}")

        def add_page_break(doc):
            try:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_break()
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to add page break: {str(e)}")

        try:
            add_page_break(ia_doc)
            add_page_break(ib_doc)
            add_page_break(ic_doc)
            add_page_break(id_doc)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to add page breaks: {str(e)}")

        try:
            composer = Composer(ia_doc)
            composer.append(ib_doc)
            composer.append(ic_doc)
            composer.append(id_doc)
            composer.append(ie_doc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to merge documents: {str(e)}")

        try:
            output = io.BytesIO()
            composer.save(output)
            output.seek(0)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save merged document: {str(e)}")

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error in merge_documents_part_i_all: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/merge-documents-part-ii")
async def merge_documents_part_ii(
    part_ii_a: UploadFile = File(...),
    part_ii_b: UploadFile = File(...),
    part_ii_c: UploadFile = File(...),
    part_ii_d: UploadFile = File(...),
):
    try:
        try:
            ii_a_bytes = io.BytesIO(await part_ii_a.read())
            ii_b_bytes = io.BytesIO(await part_ii_b.read())
            ii_c_bytes = io.BytesIO(await part_ii_c.read())
            ii_d_bytes = io.BytesIO(await part_ii_d.read())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read uploaded files: {str(e)}")

        try:
            ii_a_doc = Document(ii_a_bytes)
            ii_b_doc = Document(ii_b_bytes)
            ii_c_doc = Document(ii_c_bytes)
            ii_d_doc = Document(ii_d_bytes)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX files: {str(e)}")

        def add_page_break(doc):
            try:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_break()
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to add page break: {str(e)}")

        try:
            add_page_break(ii_a_doc)
            add_page_break(ii_b_doc)
            add_page_break(ii_c_doc)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to add page breaks: {str(e)}")

        try:
            composer = Composer(ii_a_doc)
            composer.append(ii_b_doc)
            composer.append(ii_c_doc)
            composer.append(ii_d_doc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to merge documents: {str(e)}")
        
        try:
            output = io.BytesIO()
            composer.save(output)
            output.seek(0)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save merged document: {str(e)}")

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error in merge_documents_part_ii: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/merge-documents-part-iii")
async def merge_documents_part_iii(
    part_iii_a: UploadFile = File(...),
    part_iii_b: UploadFile = File(...),
    part_iii_c: UploadFile = File(...),
):
    try:
        try:
            iii_a_bytes = io.BytesIO(await part_iii_a.read())
            iii_b_bytes = io.BytesIO(await part_iii_b.read())
            iii_c_bytes = io.BytesIO(await part_iii_c.read())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read uploaded files: {str(e)}")

        try:
            iii_a_doc = Document(iii_a_bytes)
            iii_b_doc = Document(iii_b_bytes)
            iii_c_doc = Document(iii_c_bytes)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX files: {str(e)}")

        def add_page_break(doc):
            try:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_break()
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to add page break: {str(e)}")

        try:
            add_page_break(iii_a_doc)
            add_page_break(iii_b_doc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to add page breaks: {str(e)}")

        try:
            composer = Composer(iii_a_doc)
            composer.append(iii_b_doc)
            composer.append(iii_c_doc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to merge documents: {str(e)}")

        try:
            output = io.BytesIO()
            composer.save(output)
            output.seek(0)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save merged document: {str(e)}")

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error in merge_documents_part_iii: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")
    
@app.post("/merge-documents-part-iv")
async def merge_documents_part_iv(
    part_iv_a: UploadFile = File(...),
    part_iv_b: UploadFile = File(...),
):
    try:
        try:
            iv_a_bytes = io.BytesIO(await part_iv_a.read())
            iv_b_bytes = io.BytesIO(await part_iv_b.read())
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read uploaded files: {str(e)}")

        try:
            iv_a_doc = Document(iv_a_bytes)
            iv_b_doc = Document(iv_b_bytes)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX files: {str(e)}")

        def add_page_break(doc):
            try:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_break()
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Failed to add page break: {str(e)}")

        try:
            add_page_break(iv_a_doc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to add page breaks: {str(e)}")

        try:
            composer = Composer(iv_a_doc)
            composer.append(iv_b_doc)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to merge documents: {str(e)}")

        try:
            output = io.BytesIO()
            composer.save(output)
            output.seek(0)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save merged document: {str(e)}")

        return StreamingResponse(
            output,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error in merge_documents_part_iv: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")
    
@app.post("/generate-ia-docx/")
async def generate_ia_docx_endpoint(request: Request):
    data = await request.json()
    year_range = request.headers.get('yearrange', '')
    logger.debug(f"[IA DOCX] Received yearRange from headers: '{year_range}'")
    replacements = {
        "${documentName}": data.get('documentName', ''),
        "${legalBasis}": data.get('legalBasis', ''),
        "${visionStatement}": data.get('visionStatement', ''),
        "${missionStatement}": data.get('missionStatement', ''),
        "${framework}": data.get('framework', ''),
        "${pillar}": data.get('pillar', ''),
        "${yearRange}": year_range,
    }
    import tempfile, uuid, os
    temp_dir = tempfile.gettempdir()
    filename = f"ia_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(temp_dir, filename)
    template_path = 'assets/a.docx'
    fill_placeholders_and_bullets(template_path, output_path, replacements)
    from fastapi.responses import FileResponse
    return FileResponse(output_path, filename="document.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
@app.post("/generate-iib-docx/")
async def generate_iib_docx_endpoint(request: Request):
    data = await request.json()
    systems = data.get('systems', [])  
    year_range = data.get('yearRange', '')  
    logger.debug(f"[IIB DOCX] Received yearRange: '{year_range}'")
    
    import tempfile, uuid, os
    temp_dir = tempfile.gettempdir()
    filename = f"iib_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(temp_dir, filename)
    
    template_path = 'assets/II_b.docx'
    if year_range:
        temp_template_path = os.path.join(temp_dir, f"temp_iib_template_{uuid.uuid4().hex}.docx")
        replacements = {"${yearRange}": year_range}
        fill_placeholders_and_bullets(template_path, temp_template_path, replacements)
        template_path = temp_template_path
    
    from server.tables import create_iib_docx
    create_iib_docx(systems, output_path, template_path)
    
    if year_range and os.path.exists(temp_template_path):
        try:
            os.remove(temp_template_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary template file: {e}")
    
    from fastapi.responses import FileResponse
    return FileResponse(output_path, filename="document.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
@app.post("/generate-iic-docx/")
async def generate_iic_docx_endpoint(request: Request):
    data = await request.json()
    databases = data.get('databases', data)  
    year_range = data.get('yearRange', '')  
    logger.debug(f"[IIC DOCX] Received yearRange: '{year_range}'")
    
    import tempfile, uuid, os
    temp_dir = tempfile.gettempdir()
    filename = f"iic_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(temp_dir, filename)
    
    template_path = 'assets/II_c.docx'
    if year_range:
        temp_template_path = os.path.join(temp_dir, f"temp_iic_template_{uuid.uuid4().hex}.docx")
        replacements = {"${yearRange}": year_range}
        fill_placeholders_and_bullets(template_path, temp_template_path, replacements)
        template_path = temp_template_path
    
    from server.tables import create_iic_docx
    create_iic_docx(databases, output_path, template_path)
    
    if year_range and os.path.exists(temp_template_path):
        try:
            os.remove(temp_template_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary template file: {e}")
    
    from fastapi.responses import FileResponse
    return FileResponse(output_path, filename="document.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
@app.post("/generate-iii-a-docx/")
async def generate_iii_a_docx_endpoint(request: Request):
    projects = await request.json()
    year_range = request.headers.get('yearrange', '')
    logger.debug(f"[IIIA DOCX] Received yearRange: '{year_range}'")
    
    import tempfile, uuid, os
    temp_dir = tempfile.gettempdir()
    filename = f"iii_a_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(temp_dir, filename)
    
    template_path = 'assets/III_a.docx'
    if year_range:
        temp_template_path = os.path.join(temp_dir, f"temp_iii_a_template_{uuid.uuid4().hex}.docx")
        replacements = {"${yearRange}": year_range}
        fill_placeholders_and_bullets(template_path, temp_template_path, replacements)
        template_path = temp_template_path
    
    from server.tables import create_iii_a_docx
    create_iii_a_docx(projects, output_path, template_path)
    
    if year_range and os.path.exists(temp_template_path):
        try:
            os.remove(temp_template_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary template file: {e}")
    
    from fastapi.responses import FileResponse
    return FileResponse(output_path, filename="document.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/generate-iii-b-docx/")
async def generate_iii_b_docx_endpoint(request: Request):
    projects = await request.json()
    year_range = request.headers.get('yearrange', '')
    logger.debug(f"[IIIB DOCX] Received yearRange: '{year_range}'")
    
    import tempfile, uuid, os
    temp_dir = tempfile.gettempdir()
    filename = f"iii_b_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(temp_dir, filename)
    
    template_path = 'assets/III_b.docx'
    if year_range:
        temp_template_path = os.path.join(temp_dir, f"temp_iii_b_template_{uuid.uuid4().hex}.docx")
        replacements = {"${yearRange}": year_range}
        fill_placeholders_and_bullets(template_path, temp_template_path, replacements)
        template_path = temp_template_path
    
    from server.tables import create_iii_b_docx
    create_iii_b_docx(projects, output_path, template_path)
    
    if year_range and os.path.exists(temp_template_path):
        try:
            os.remove(temp_template_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary template file: {e}")
    
    from fastapi.responses import FileResponse
    return FileResponse(output_path, filename="document.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/convert-docx")
async def convert_docx(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No selected file")
    try:
        contents = await file.read()
        with open("temp_upload.docx", "wb") as f:
            f.write(contents)
        with open("temp_upload.docx", "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
        os.remove("temp_upload.docx")
        return {"html": html}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to convert DOCX: {str(e)}")

@app.get("/")
async def root():
    return {"message": "Document Merge Server is running"}

def fill_placeholders_and_bullets(template_path, output_path, replacements):
    from docx import Document
    from docx.shared import Pt
    doc = Document(template_path)

    for para in doc.paragraphs:
        for ph, val in replacements.items():
            if ph in para.text:
                if ph == "${framework}":
                    para.clear()
                    run = para.add_run(val)
                    run.bold = True
                    run.font.name = 'Palatino Linotype'
                    run.font.size = Pt(12)
                elif ph in ["${visionStatement}", "${missionStatement}"]:
                    para.clear()
                    run = para.add_run(f'"{val}"')
                    run.italic = True
                    run.font.name = 'Palatino Linotype'
                    run.font.size = Pt(12)
                else:
                    para.text = para.text.replace(ph, val)
                    for run in para.runs:
                        run.font.name = 'Palatino Linotype'
                        run.font.size = Pt(12)
    
    for para in doc.paragraphs:
        replace_placeholder_in_paragraph(para, "${pillar}", replacements.get("${pillar}", ""))
        for run in para.runs:
            run.font.name = 'Palatino Linotype'
            run.font.size = Pt(12)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for para in header.paragraphs:
                for ph, val in replacements.items():
                    if ph in para.text:
                        para.text = para.text.replace(ph, val)
                        for run in para.runs:
                            run.font.name = 'Palatino Linotype'
                            run.font.size = Pt(14)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for para in footer.paragraphs:
                for ph, val in replacements.items():
                    if ph in para.text:
                        para.text = para.text.replace(ph, val)
                        for run in para.runs:
                            run.font.name = 'Palatino Linotype'
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(0, 0, 0)
    doc.save(output_path)

def replace_placeholder_in_paragraph(para, placeholder, replacement):
    full_text = ''.join(run.text for run in para.runs)
    if placeholder in full_text:
        new_text = full_text.replace(placeholder, replacement)
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)

@app.post("/generate-ib-docx/")
async def generate_ib_docx_endpoint(request: Request):
    try:
        data = await request.json()
        year_range = request.headers.get('yearrange', '') or data.get('yearRange', '')
        logger.debug(f"[IB DOCX] Received yearRange: '{year_range}'")
        
        current_date = data.get('currDate', '')
        if not current_date:
            current_date = datetime.now().strftime('%B %d, %Y')
        data['currDate'] = current_date
        
        processed_data = process_part_ib_data(data)
        processed_data['yearRange'] = year_range
        logger.debug(f"Processed data: {processed_data}")
        
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'assets', 'b.docx')
        logger.debug(f"Looking for template at: {template_path}")
        
        if not os.path.exists(template_path):
            error_msg = f"Template file not found at {template_path}"
            logger.error(error_msg)
            raise HTTPException(status_code=404, detail=error_msg)

        try:
            docx_bytes = generate_part_ib_docx(processed_data, template_path)
            logger.debug("Document generated successfully")
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
        
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=part_ib.docx"}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in generate_ib_docx_endpoint: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.post("/generate-iiic-docx/")
async def generate_iiic_docx_endpoint(request: Request):
    try:
        data = await request.json()
        year_range = request.headers.get('yearrange', '')
        logger.debug(f"[IIIC DOCX] Received yearRange: '{year_range}'")
        logger.debug(f"[IIIC DOCX] Received data: {data}")
        
        import tempfile, uuid, os
        temp_dir = tempfile.gettempdir()
        filename = f"iiic_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(temp_dir, filename)
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'assets', 'III_c.docx')
        
        logger.debug(f"[IIIC DOCX] Template path: {template_path}")
        if not os.path.exists(template_path):
            error_msg = f"Template file not found at {template_path}"
            logger.error(error_msg)
            raise HTTPException(status_code=404, detail=error_msg)
        
        if year_range:
            temp_template_path = os.path.join(temp_dir, f"temp_iiic_template_{uuid.uuid4().hex}.docx")
            replacements = {"${yearRange}": year_range}
            fill_placeholders_and_bullets(template_path, temp_template_path, replacements)
            template_path = temp_template_path
            logger.debug(f"[IIIC DOCX] Using temp template: {template_path}")
        
        doc = Document(template_path)
        logger.debug(f"[IIIC DOCX] Document loaded successfully")
        
        # Handle the data structure from frontend
        logframes = data.get('logframes', [])
        if not logframes:
            # Fallback for single logframe structure
            logframes = [data]
        
        logger.debug(f"[IIIC DOCX] Processing {len(logframes)} logframes")
        
        # Process each logframe
        for i, logframe in enumerate(logframes):
            logger.debug(f"[IIIC DOCX] Processing logframe {i+1}: {logframe}")
            if i > 0:
                # Add spacing between multiple logframes
                doc.add_paragraph()
                doc.add_paragraph()
            
            # Create the logframe table for this project (don't pass output_path when using existing doc)
            create_iiic_logframe_table(logframe, None, doc=doc)
        
        # Save the final document
        doc.save(output_path)
        logger.debug(f"[IIIC DOCX] Document saved to: {output_path}")
        
        if year_range and os.path.exists(temp_template_path):
            try:
                os.remove(temp_template_path)
            except Exception as e:
                logger.warning(f"Failed to remove temporary template file: {e}")
        
        from fastapi.responses import FileResponse
        return FileResponse(output_path, filename="document.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in generate_iiic_docx_endpoint: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--port", type=int, default=8000, help="Port to run the server on")
    args = parser.parse_args()
    uvicorn.run(app, host="0.0.0.0", port=args.port) 