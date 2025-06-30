import os
import io
import base64
import tempfile
import traceback
import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import qn

# Set up logger
logger = logging.getLogger(__name__)

def replace_first_picture_content_control(doc, image_path):
    """
    Replace the first picture content control in the document with the given image.
    """
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for sdt in doc.element.iterfind('.//%ssdt' % ns):
        sdt_content = sdt.find(f'{ns}sdtContent')
        if sdt_content is not None:
            found_pic = False
            for descendant in sdt_content.iter():
                if descendant.tag in (f'{ns}drawing', f'{ns}pict'):
                    found_pic = True
                    break
            if found_pic:
                for child in list(sdt_content):
                    sdt_content.remove(child)
                from docx.text.paragraph import Paragraph
                from docx.oxml import OxmlElement
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                p.append(r)
                sdt_content.append(p)
                para = Paragraph(p, doc)
                run = para.add_run()
                run.add_picture(image_path, width=Inches(6))
                break

def generate_part_ib_docx(data: Dict[str, Any], template_path: str) -> bytes:
    """
    Generate a DOCX file for Part IB using the provided data and template.
    Inserts the image at the paragraph containing the ${organizationalStructure} placeholder.
    
    Args:
        data: Dictionary containing the form data
        template_path: Path to the template DOCX file
        
    Returns:
        bytes: The generated DOCX file as bytes
    """
    try:
        if not os.path.exists(template_path):
            error_msg = f"Template file not found at {template_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)
            
        try:
            doc = Document(template_path)
            logger.debug("Template loaded successfully")
        except Exception as e:
            error_msg = f"Error loading template: {str(e)}"
            logger.error(error_msg)
            raise
            
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key != 'organizationalStructure':
                    placeholder = f'${{{key}}}'
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if key != 'organizationalStructure':
                            placeholder = f'${{{key}}}'
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
        
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                for para in header.paragraphs:
                    for key, value in data.items():
                        placeholder = f'${{{key}}}'
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))
                            for run in para.runs:
                                run.font.name = 'Palatino Linotype'
                                run.font.size = Pt(14)
                                run.bold = True
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                for para in footer.paragraphs:
                    for key, value in data.items():
                        placeholder = f'${{{key}}}'
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))
        
        if 'organizationalStructure' in data and data['organizationalStructure']:
            try:
                image_data = base64.b64decode(data['organizationalStructure'])
                logger.debug(f"Decoded image size: {len(image_data)} bytes")
                
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                    temp_img.write(image_data)
                    temp_img_path = temp_img.name

                found = False
                for i, paragraph in enumerate(doc.paragraphs):
                    if '${organizationalStructure}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('${organizationalStructure}', '')
                        run = paragraph.add_run()
                        run.add_picture(temp_img_path, width=Inches(8.29), height=Inches(5.27))
                        found = True
                        break
                if not found:
                    logger.warning("Image placeholder not found in document.")
                
                os.unlink(temp_img_path)
            except Exception as e:
                error_msg = f"Error processing image: {str(e)}"
                logger.error(error_msg)
                raise
        else:
            logger.warning("No organizational structure image provided")
        
        try:
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            return docx_bytes.getvalue()
        except Exception as e:
            error_msg = f"Error saving document: {str(e)}"
            logger.error(error_msg)
            raise
            
    except Exception as e:
        error_msg = f"Error in generate_part_ib_docx: {str(e)}"
        logger.error(error_msg)
        raise

def format_number(value: str) -> str:
    """
    Format a number string with commas.
    
    Args:
        value: The number string to format
    
    Returns:
        str: Formatted number string or original value if not a number
    """
    if not value or not value.strip():
        return '0'
    value = value.strip().replace(',', '')
    if not value.replace('-', '').replace('.', '').isdigit():
        return value
        
    try:
        num = int(float(value))
        return f"{num:,}"
    except (ValueError, TypeError) as e:
        return value

def process_part_ib_data(data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Process the raw form data for Part IB.
    
    Args:
        data: Raw form data from the request
    
    Returns:
        Dict[str, Any]: Processed data ready for document generation
    """
    try:
        numeric_fields = [
            'mooe', 'co', 'total', 'nicthsProjectCost', 'hsdvProjectCost', 'hecsProjectCost',
            'totalEmployees', 'regionalOffices', 'provincialOffices',
            'coPlantilaPositions', 'coVacant', 'coFilledPlantilaPositions', 
            'coFilledPhysicalPositions', 'coCosws', 'coContractual', 'coTotal',
            'foPlantilaPositions', 'foVacant', 'foFilledPlantilaPositions',
            'foFilledPhysicalPositions', 'foCosws', 'foContractual', 'foTotal'
        ]
        
        processed_data = data.copy()
        
        for field in numeric_fields:
            if field in processed_data:
                processed_data[field] = format_number(processed_data[field])
        
        return processed_data
    except Exception as e:
        error_msg = f"Error processing Part IB data: {str(e)}"
        logger.error(error_msg)
        raise

def generate_ib_docx(data):
    try:
        template_path = os.path.join(os.path.dirname(__file__), 'templates', 'part_ib_template.docx')
        doc = Document(template_path)

        replacements = {
            '${plannerName}': data.get('plannerName', ''),
            '${plantillaPosition}': data.get('plantillaPosition', ''),
            '${organizationalUnit}': data.get('organizationalUnit', ''),
            '${emailAddress}': data.get('emailAddress', ''),
            '${contactNumbers}': data.get('contactNumbers', ''),
            '${mooe}': data.get('mooe', ''),
            '${co}': data.get('co', ''),
            '${total}': data.get('total', ''),
            '${totalEmployees}': data.get('totalEmployees', ''),
            '${regionalOffices}': data.get('regionalOffices', ''),
            '${provincialOffices}': data.get('provincialOffices', ''),
            '${otherOffices}': data.get('otherOffices', ''),
            '${coPlantilaPositions}': data.get('coPlantilaPositions', ''),
            '${coVacant}': data.get('coVacant', ''),
            '${coFilledPlantilaPositions}': data.get('coFilledPlantilaPositions', ''),
            '${coFilledPhysicalPositions}': data.get('coFilledPhysicalPositions', ''),
            '${coCosws}': data.get('coCosws', ''),
            '${coContractual}': data.get('coContractual', ''),
            '${coTotal}': data.get('coTotal', ''),
            '${foPlantilaPositions}': data.get('foPlantilaPositions', ''),
            '${foVacant}': data.get('foVacant', ''),
            '${foFilledPlantilaPositions}': data.get('foFilledPlantilaPositions', ''),
            '${foFilledPhysicalPositions}': data.get('foFilledPhysicalPositions', ''),
            '${foCosws}': data.get('foCosws', ''),
            '${foContractual}': data.get('foContractual', ''),
            '${foTotal}': data.get('foTotal', ''),
            '${otrFund}': data.get('otrFund', ''),
            '${currDate}': data.get('currDate', ''),
        }

        bold_keys = {'${total}', '${coTotal}', '${foTotal}'}
        def replace_placeholder_with_bold(paragraph, key, value):
            full_text = ''.join(run.text for run in paragraph.runs)
            if key not in full_text:
                return
            parts = full_text.split(key)
            for run in paragraph.runs:
                run.text = ''
            for i, part in enumerate(parts):
                if part:
                    paragraph.add_run(part)
                if i < len(parts) - 1:
                    bold_run = paragraph.add_run(value)
                    bold_run.bold = True

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in bold_keys and key in paragraph.text:
                    replace_placeholder_with_bold(paragraph, key, value)
                elif key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        for para in cell.paragraphs:
                            if key in bold_keys and key in para.text:
                                replace_placeholder_with_bold(para, key, value)
                            elif key in para.text:
                                para.text = para.text.replace(key, value)

        temp_path = os.path.join(os.path.dirname(__file__), 'temp', 'part_ib_output.docx')
        os.makedirs(os.path.dirname(temp_path), exist_ok=True)
        doc.save(temp_path)

        return temp_path
    except Exception as e:
        error_msg = f"Error generating Part IB document: {str(e)}"
        logger.error(error_msg)
        raise 